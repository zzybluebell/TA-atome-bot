from io import BytesIO
from langchain_core.documents import Document


class DocumentReader:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    def __init__(
        self,
        chunk_size: int = 1200,
        chunk_overlap: int = 200,
        max_pdf_pages: int = 150,
        max_text_characters: int = 1_000_000,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_pdf_pages = max_pdf_pages
        self.max_text_characters = max_text_characters

    def read_bytes(self, file_name: str, content: bytes) -> list[Document]:
        extension = self._get_extension(file_name)
        text = self._read_text_by_extension(extension, content)
        chunks = self._split_text(text)
        return [
            Document(
                page_content=chunk,
                metadata={
                    "source": file_name,
                    "file_type": extension.lstrip("."),
                    "chunk_index": index,
                },
            )
            for index, chunk in enumerate(chunks)
        ]

    def _get_extension(self, file_name: str) -> str:
        dot_index = file_name.rfind(".")
        extension = file_name[dot_index:].lower() if dot_index >= 0 else ""
        if extension not in self.SUPPORTED_EXTENSIONS:
            supported = ", ".join(sorted(self.SUPPORTED_EXTENSIONS))
            raise ValueError(f"Unsupported file type: {extension or '<none>'}. Supported: {supported}")
        return extension

    def _read_text_by_extension(self, extension: str, content: bytes) -> str:
        if extension in {".txt", ".md"}:
            return self._read_text_file(content)
        if extension == ".pdf":
            return self._read_pdf(content)
        if extension == ".docx":
            return self._read_docx(content)
        raise ValueError(f"Unsupported file type: {extension}")

    def _read_text_file(self, content: bytes) -> str:
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                text = content.decode(encoding)
                if text.strip():
                    return text
            except UnicodeDecodeError:
                continue
        raise ValueError("Unable to decode text file content")

    def _read_pdf(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except Exception:
            raise ValueError("PDF support requires package: pypdf")

        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted:
            raise ValueError("Encrypted PDF is not supported")
        if len(reader.pages) > self.max_pdf_pages:
            raise ValueError(f"PDF exceeds page limit ({self.max_pdf_pages} pages)")
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if not text:
            raise ValueError("PDF file has no extractable text")
        if len(text) > self.max_text_characters:
            raise ValueError(f"PDF text exceeds character limit ({self.max_text_characters})")
        return text

    def _read_docx(self, content: bytes) -> str:
        try:
            from docx import Document as DocxDocument
        except Exception:
            raise ValueError("DOCX support requires package: python-docx")

        doc = DocxDocument(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        text = "\n".join(paragraphs).strip()
        if not text:
            raise ValueError("DOCX file has no extractable text")
        if len(text) > self.max_text_characters:
            raise ValueError(f"DOCX text exceeds character limit ({self.max_text_characters})")
        return text

    def _split_text(self, text: str) -> list[str]:
        cleaned = text.strip()
        if not cleaned:
            raise ValueError("Document content is empty")
        if len(cleaned) > self.max_text_characters:
            raise ValueError(f"Document text exceeds character limit ({self.max_text_characters})")
        if len(cleaned) <= self.chunk_size:
            return [cleaned]

        chunks: list[str] = []
        start = 0
        stride = max(1, self.chunk_size - self.chunk_overlap)
        while start < len(cleaned):
            end = min(len(cleaned), start + self.chunk_size)
            chunks.append(cleaned[start:end])
            if end == len(cleaned):
                break
            start += stride
        return chunks
